
from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from datetime import datetime
from docx.shared import Inches
import openai
import os
import fitz  # for PDF
import re

app = Flask(__name__)
CORS(app)

openai.api_key = os.getenv("OPENAI_API_KEY")

REPORT_FOLDER = os.path.join(app.root_path, 'static', 'reports')
LOGO_PATH = os.path.join(app.root_path, 'static', 'logo.png')
os.makedirs(REPORT_FOLDER, exist_ok=True)

def extract_text_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def extract_text_pdf(file):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    return "\n".join([page.get_text() for page in doc])

def extract_text(file_storage):
    filename = file_storage.filename.lower()
    if filename.endswith(".pdf"):
        return extract_text_pdf(file_storage)
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        return extract_text_docx(file_storage)
    else:
        return ""

def clean_markdown(text):
    text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
    text = text.replace('*', '')
    return text.strip()

def extract_table_data(text):
    table = []
    lines = text.strip().splitlines()
    for line in lines:
        if '|' in line:
            row = [cell.strip() for cell in line.split('|') if cell.strip()]
            if row:
                table.append(row)
    return table if len(table) >= 2 else None

def generate_section(prompt):
    try:
        print("Sending prompt to OpenAI...")
        response = openai.ChatCompletion.create(
            model="gpt-4-0125-preview",
            messages=[
                {"role": "system", "content": "You are a business consultant creating professional business plans."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1500
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        print("OpenAI API error:", e)
        return "Error generating this section."

def add_logo(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    if os.path.exists(LOGO_PATH):
        run.add_picture(LOGO_PATH, width=Inches(1.73), height=Inches(0.83))
        paragraph.alignment = 1

@app.route('/')
def home():
    return "Intro Framework Backend is Running!"

@app.route('/generate', methods=['POST'])
def generate_report():
    framework = request.files.get('framework')
    doc1 = request.files.get('doc1')
    doc2 = request.files.get('doc2')

    context = ""
    if framework: context += extract_text(framework)
    if doc1: context += "\n" + extract_text(doc1)
    if doc2: context += "\n" + extract_text(doc2)

    if not context.strip():
        return jsonify({'error': 'No valid input provided.'}), 400

    doc = Document()
    add_logo(doc)
    doc.add_heading('Business Plan', 0)
    doc.add_paragraph(f"Generated by CamTech Consulting on {datetime.now().strftime('%B %d, %Y')}")

    sections = [
        ("Executive Summary", "Write a 500-word executive summary for a business plan."),
        ("Business Description", "Describe the business, its mission, and purpose."),
        ("Market Analysis", "Analyze the market, trends, and target customer."),
        ("Organization & Management", "Describe the business structure, leadership, and team."),
        ("Products or Services", "Detail the main products or services offered."),
        ("Marketing & Sales Strategy", "Explain how the business will attract and retain customers."),
        ("Financial Overview", "Provide a high-level financial outlook including a simple table of cost/revenue."),
        ("Roadmap & Next Steps", "Outline the short-term goals and implementation plan.")
    ]

    for title, instruction in sections:
        doc.add_heading(title, level=1)
        prompt = f"{instruction}\n\nContext:\n{context}"
        gpt_response = generate_section(prompt)
        table_data = extract_table_data(gpt_response)
        if table_data:
            table = doc.add_table(rows=1, cols=len(table_data[0]))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, val in enumerate(table_data[0]):
                hdr_cells[i].text = val
            for row_data in table_data[1:]:
                row_cells = table.add_row().cells
                for i, val in enumerate(row_data):
                    if i < len(row_cells):
                        row_cells[i].text = val
        else:
            doc.add_paragraph(clean_markdown(gpt_response))

    filename = f"business_plan_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(REPORT_FOLDER, filename)
    doc.save(filepath)

    return jsonify({'download_url': f'/static/reports/{filename}'})

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
